from datetime import date
from selenium import webdriver
from selenium.webdriver.common.by import By
from sortedcontainers import SortedDict
import requests
from PIL import Image
import io
import docx
from docx.shared import Inches
from docx.enum.text import WD_BREAK
import os.path as osp

def create_doc(doc_heading:str = None):
    # create an instance of an empty word document
    doc = docx.Document()
    # Add a title to the document
    if doc_heading:
        doc.add_heading(doc_heading,0)
    return doc

def save_doc(doc, doc_name:str):
    doc_name = '_'.join([doc_name,str(date.today())])
    doc_name += '.docx'
    doc.save(osp.join('output',doc_name))
    
def clean_url(url:str):
    if url:
        return url.split('?')[0]
    raise TypeError("Not correct xiaohongshu url!")

def create_image_dict(url:str)->dict:
    driver = webdriver.Chrome()
    driver.get(clean_url(url))
    images = SortedDict()
    slides = driver.find_elements(By.XPATH,"//div[@class='swiper-wrapper']//div")
    for slide in slides:
        res_str = slide.get_attribute('style').split('"')[1]
        images[int(slide.get_attribute('data-swiper-slide-index'))] = res_str
    return images

def add_image_from_web(doc, image_url):
    # Standard dimensions for a Word document page (usually in inches)
    page_width_inch = 6.5  # This accounts for 1-inch margins on each side
    page_height_inch = 9   # This accounts for 1-inch top and bottom margins

    response = requests.get(image_url)

    if response.status_code == 200:
        image_webp = Image.open(io.BytesIO(response.content))
        # Get document page size in pixels at 72 DPI (Word's default DPI)
        page_width_px = int(page_width_inch * 72)
        page_height_px = int(page_height_inch * 72)

        # Scale the image to fit the width while maintaining aspect ratio
        image_ratio = image_webp.width / image_webp.height
        page_ratio = page_width_px / page_height_px

        if image_ratio > page_ratio:
            # Image is wider than page, scale by width
            new_width = page_width_px
            new_height = round(page_width_px / image_ratio)
        else:
            # Image is taller than page, scale by height
            new_height = page_height_px
            new_width = round(page_height_px * image_ratio)

        # Resize the image using Resampling.LANCZOS
        image_resized = image_webp.resize((new_width, new_height), Image.LANCZOS)

        # Convert to a compatible format (PNG)
        image_stream = io.BytesIO()
        image_resized.save(image_stream, format='PNG')

        # Seek to the beginning of the stream
        image_stream.seek(0)

        # Add the image to the document with a width of 6.5 inches
        doc.add_picture(image_stream, width=Inches(page_width_inch))

        # Add a page break
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        print('The image has been added and scaled to fit the page in the doc')
    else:
        print(f"Failed to retrieve the image. Status code: {response.status_code}")